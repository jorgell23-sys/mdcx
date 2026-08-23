# Copyright 2026 Jorge Ellena G.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Conversion engines.

Docling understands layout, heading hierarchy and tables. Native extractors
do not interpret structure but are deterministic and depend on no model.

No engine uses a language model to transcribe: a generative model may
paraphrase, and the requirement is fidelity to the source.
"""
from __future__ import annotations

import os as _os
import html
import re
from pathlib import Path

_CONVERTERS: dict = {}

_DEVICE_CACHE: dict = {}

# How many processes may use the card at once. Video memory is what runs out
# first here: a process holding the table models with a full batch takes around
# 1,384 MiB, so eight of them ask for 11 GB of a card that has 6. Measured, the
# run does not fail -- it stalls, at 100% utilisation and 5,714 MiB of 6,144,
# with nine documents of twelve done after eight minutes and no further
# progress. With three processes allowed on the card the same twelve finish in
# 72 seconds.
#
# The lane used to be what bounded this, by being small and being the only lane
# whose workers were given the engines at all. That made the lane decide two
# separate things -- which documents may use the card, and which may use the
# structured engine -- and the second was never what anyone wanted: it cost the
# CPU lane a third of its table rows and every list item.
#
# So the bound is stated here instead, where the card is actually used, and the
# lane goes back to being about scheduling. Set per worker; where it is not set,
# as in a single process or a test, there is nothing to contend for and nothing
# to wait on.
_GPU_GATE = None


def set_gpu_gate(gate) -> None:
    """Give this worker the gate that limits how many processes use the card."""
    global _GPU_GATE
    _GPU_GATE = gate


class _Turn:
    """The right to use the card, released as soon as the model is done with it.

    Held around the model call rather than around the document, so a process
    that spends most of its time extracting text is not occupying a place on
    the card while it does.
    """

    def __enter__(self):
        if _GPU_GATE is not None:
            _GPU_GATE.acquire()
        return self

    def __exit__(self, *excepcion) -> bool:
        if _GPU_GATE is not None:
            _GPU_GATE.release()
        return False


def gpu_turn() -> "_Turn":
    return _Turn()

def gpu_available() -> bool:
    if "gpu" not in _DEVICE_CACHE:
        try:
            import torch

            _DEVICE_CACHE["gpu"] = bool(torch.cuda.is_available())
            _DEVICE_CACHE["name"] = (
                torch.cuda.get_device_name(0) if _DEVICE_CACHE["gpu"] else "CPU"
            )
        except Exception:
            _DEVICE_CACHE["gpu"] = False
            _DEVICE_CACHE["name"] = "CPU"
    return _DEVICE_CACHE["gpu"]

def device_name() -> str:
    gpu_available()
    return _DEVICE_CACHE.get("name", "CPU")

def _accelerator_options(threads: int | None = None):
    """Configure Docling to use CUDA when available, falling back cleanly to CPU.

    The thread count is the budget this worker was started with, not a constant.
    Four apiece was the constant, and eight workers asking for four threads each
    is thirty-two threads on a machine of twelve: the share the process cap was
    careful to leave free was spent again inside the pool, and the machine
    stopped being usable while a library converted.
    """
    import os as _os

    if threads is None:
        try:
            threads = max(1, int(_os.environ.get("OMP_NUM_THREADS", "4")))
        except ValueError:
            threads = 4
    try:
        from docling.datamodel.pipeline_options import AcceleratorDevice, AcceleratorOptions

        device = AcceleratorDevice.CUDA if gpu_available() else AcceleratorDevice.AUTO
        return AcceleratorOptions(num_threads=threads, device=device)
    except Exception:
        return None

def local_artifacts_path():
    """Folder holding already downloaded models, if present."""
    from pathlib import Path as _P
    import os as _os

    candidatas = []
    env = _os.environ.get("DOCLING_ARTIFACTS_PATH")
    if env:
        candidatas.append(_P(env))
    candidatas.append(_P(__file__).resolve().parent.parent / "models")
    candidatas.append(_P.home() / ".cache" / "docling" / "models")
    for c in candidatas:
        if c.is_dir() and any(c.iterdir()):
            return c
    return None

def docling_available() -> bool:
    try:
        import docling  # noqa: F401

        return True
    except Exception:
        return False

def _get_docling_converter(ocr: bool):
    """Crea (y cachea) un DocumentConverter. Cachear importa: construirlo carga modelos."""
    key = f"ocr={ocr}"
    if key in _CONVERTERS:
        return _CONVERTERS[key]

    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions

    popts = PdfPipelineOptions()
    popts.do_ocr = ocr
    popts.do_table_structure = True
    accel = _accelerator_options()
    if accel is not None:
        popts.accelerator_options = accel

    artefactos = local_artifacts_path()
    if artefactos is not None:
        popts.artifacts_path = str(artefactos)
    try:
        from docling.datamodel.pipeline_options import TableFormerMode

        modo = _os.environ.get("PDFTOMD_TABLAS", "rapido").strip().lower()
        popts.table_structure_options.mode = (
            TableFormerMode.ACCURATE if modo in ("exacto", "accurate")
            else TableFormerMode.FAST)
        popts.table_structure_options.do_cell_matching = True
    except Exception:
        pass
    if ocr:
        try:
            from docling.datamodel.pipeline_options import RapidOcrOptions

            popts.ocr_options = RapidOcrOptions(
                lang=["english"],
                force_full_page_ocr=True,
                scale=4.0,
            )
        except Exception:
            try:
                popts.ocr_options.force_full_page_ocr = True
            except Exception:
                pass

    conv = DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=popts)}
    )
    _CONVERTERS[key] = conv
    return conv

def _full_export_kwargs() -> dict:
    """Force Docling to export the whole document, not only the body."""
    kwargs: dict = {}
    try:
        from docling_core.types.doc.document import ContentLayer

        kwargs["included_content_layers"] = set(ContentLayer)
    except Exception:
        pass
    try:
        from docling_core.types.doc.labels import DocItemLabel

        kwargs["labels"] = set(DocItemLabel)
    except Exception:
        pass
    return kwargs

def docling_convert(path: Path, ocr: bool = False) -> tuple[str, dict, dict | None]:
    """Devuelve (markdown, metadatos, documento_json_sin_perdida)."""
    conv = _get_docling_converter(ocr)
    # Layout analysis is the other thing here that reserves the card, and the
    # heaviest: a document without a text layer is read page by page by a model.
    with gpu_turn():
        result = conv.convert(str(path))
    doc = result.document
    try:
        md = doc.export_to_markdown(**_full_export_kwargs())
    except Exception:
        md = doc.export_to_markdown()
    try:
        lossless = doc.export_to_dict()
    except Exception:
        lossless = None
    meta = {"engine": "docling", "ocr": ocr}
    try:
        meta["pages"] = len(doc.pages)
    except Exception:
        pass
    return md, meta, lossless

# How much of a page's text a table must account for before the table stands
# in for the page. Below it the prose is kept as well: the table is then a
# small part of the page and dropping the rest loses more than it saves.
TABLE_DOMINATES = 0.6

def _extract_pages(path: Path, headings: dict | None = None) -> tuple[list[str], dict]:
    """Each page as Markdown, and which pages this engine could not settle.

    Tables are located from what is drawn on the page and from the caption the
    author wrote, never from text alignment alone; `tables` says why.

    A page that announces a table and yields none is reported rather than
    quietly accepted. It is the one thing this engine knows it may have got
    wrong, and naming it is what lets a caller send those pages -- and only
    those -- to something that can read them.
    """
    from . import pdf as _pdf
    from . import tables as _tables

    # The section titles the author already wrote into the document. Extracting
    # prose gives no headings at all, and a chapter without them is a wall of
    # text: nothing to cite, and the titles themselves are among the best
    # answers a search can return. Where the caller knows the document these
    # pages were cut from, it passes them in, because cutting pages leaves the
    # bookmarks behind.
    if headings is None:
        headings = {}
        for level, title, page in _pdf.outline(path):
            headings.setdefault(page, []).append((level, title))

    doc = _pdf.open_document(path)
    pages: list[str] = []
    found = 0
    announced = 0
    titled = 0
    unresolved: list[int] = []
    try:
        for index, page in enumerate(doc):
            table = None
            says_so = False
            ruled = False
            try:
                textpage = page.get_textpage()
                seen = _tables.examine(page, textpage)
                table, says_so, ruled = seen["table"], seen["announced"], seen["ruled"]
            except Exception:  # noqa: BLE001
                table = None        # a page that resists is still prose
            if says_so:
                announced += 1

            parts = [f"\n<!-- page {index + 1} -->\n"]
            for level, title in headings.get(index + 1, []):
                parts.append("#" * min(max(level, 1), 6) + " " + title)
                titled += 1
            prose = [t for t in _pdf.page_paragraphs_fast(page) if t]
            if table:
                found += 1
                written = sum(len(t) for t in prose)
                # A page that is a table should be the table: read as running
                # text its cells arrive in reading order, which is worse than
                # not having them. But on most pages the table is the smaller
                # part, and replacing the page with it throws away the prose
                # around it -- measured here, a table of 237 characters was
                # costing 1,530 characters of text on the same page.
                if written and len(table) < written * TABLE_DOMINATES:
                    parts.extend(prose)
                parts.append(table)
            else:
                # A page is worth a second look either because the author said
                # there was a table on it, or because it is ruled the way a
                # table is ruled. The label is the stronger signal but it is
                # not a census: a book that calls a screenshot of a spreadsheet
                # "Figure 4.2" announces no tables at all, and counting zero as
                # "there are none" leaves the whole document unexamined. The
                # rules are drawn whatever the author chose to call it.
                if says_so or ruled:
                    unresolved.append(index)
                parts.extend(prose)
            pages.append("\n\n".join(parts))
        meta = {"engine": "pypdfium2", "pages": len(doc), "tables": found,
                "tables_announced": announced, "unresolved": unresolved,
                "headings": titled}
    finally:
        doc.close()
    return pages, meta

# The extraction of the document being converted, kept only until another one
# starts.
#
# Both cheap engines begin here, and the pipeline offers them one after the
# other: the native attempt extracts the document, and the hybrid -- which is
# the native engine plus a second look at the pages it could not settle -- used
# to extract it again. Same path, same headings, nothing in between, so by
# construction the same result. Measured at 3.0 ms a page, on every document
# that reaches the hybrid: 208 of 292 chapters in one run, about a tenth of it.
#
# One entry, because the repeat is immediate. Keeping more would hold whole
# books in memory for a saving that has already been taken.
_LAST_EXTRACTION: dict = {}


def _native_pages(path: Path, headings: dict | None = None) -> tuple[list[str], dict]:
    """The extraction, done once per document however many engines ask for it."""
    if (_LAST_EXTRACTION.get("path") == str(path)
            and _LAST_EXTRACTION.get("headings") == headings):
        pages, meta = _LAST_EXTRACTION["result"]
    else:
        pages, meta = _extract_pages(path, headings)
        _LAST_EXTRACTION.clear()
        _LAST_EXTRACTION.update(path=str(path), headings=headings,
                                result=(pages, meta))

    # A copy, because a caller edits what it is given: _read_shapes replaces,
    # in place, each page it recovered a table from. Sharing the list would
    # hand the next engine a document the previous one had already rewritten,
    # which is the same class of fault as sharing it between packages.
    copia = dict(meta)
    if isinstance(copia.get("unresolved"), list):
        copia["unresolved"] = list(copia["unresolved"])
    return list(pages), copia


def native_pdf(path: Path, headings: dict | None = None) -> tuple[str, dict, None]:
    """PDF to Markdown by layout blocks, preserving reading order and page separation."""
    pages, meta = _native_pages(path, headings)
    return "\n\n".join(pages), meta, None

def _read_shapes(path: Path, pages: list[str], meta: dict,
                 pending: list[int]) -> tuple[str, dict, None]:
    """Recover the tables of the pending pages by reading their shape.

    Only the shape: where the rows and the columns run. The words come from the
    text layer as everywhere else, so nothing is transcribed and nothing can be
    invented into a cell -- and a page with no text layer gains nothing here,
    which is what OCR is for.

    The table replaces the page it was found on, the same way the cheap path
    replaces a page its table dominates. A page that yields no table keeps the
    prose already extracted from it.
    """
    from . import pdf as _pdf
    from . import tatr as _tatr

    recovered = 0
    document = _pdf.open_document(path)
    try:
        targets = [document[index] for index in pending]
        textpages = [page.get_textpage() for page in targets]
        try:
            with gpu_turn():
                found = _tatr.tables_on_pages(targets, textpages)
        except Exception:  # noqa: BLE001
            found = [None] * len(targets)
        for position, (index, table) in enumerate(zip(pending, found)):
            if not table:
                continue
            # The prose is kept whatever the table covers, which is not what the
            # cheap path does. There the table is drawn and takes up a part of
            # the page, so replacing the page where it dominates costs nothing;
            # here the table was found by a model that frames it generously, and
            # dropping the page around it dropped text.
            #
            # Measured over chapters of a textbook: keeping the prose took
            # coverage from 0.992 to 1.000 with exactly the same rows of table,
            # for two and a half per cent more characters. The duplication is
            # real -- a cell's words appear in the row and in the paragraph --
            # and it is the cheaper of the two mistakes.
            prose = [t for t in _pdf.page_paragraphs_fast(targets[position]) if t]
            parts = [f"\n<!-- page {index + 1} -->\n"]
            parts.extend(prose)
            parts.append(table)
            pages[index] = "\n\n".join(parts)
            recovered += 1
    finally:
        document.close()

    meta = dict(meta)
    meta["engine"] = "pypdfium2+tatr" if recovered else "pypdfium2"
    meta["pages_read_by_model"] = len(pending)
    meta["tables_recovered"] = recovered
    return "\n\n".join(p for p in pages if p), meta, None

def hybrid_pdf(path: Path, headings: dict | None = None) -> tuple[str, dict, None]:
    """Extract the whole document, and read with the model only what needs it.

    Deciding the engine once per document is what makes the expensive one
    expensive: a book with tables on a tenth of its pages pays for layout
    analysis on the other nine tenths. The decision belongs to the page.

    So every page is extracted natively, and the few that announce a table the
    cheap path could not produce are handed to the model one at a time. Where
    that is a handful of pages in a book, the cost of the model is charged on a
    handful of pages rather than on all of them.
    """
    import tempfile

    from . import pdf as _pdf
    from . import tatr as _tatr

    pages, meta = _native_pages(path, headings)
    pending = list(meta.get("unresolved") or [])
    if not pending:
        return "\n\n".join(pages), meta, None

    # A model that reads the shape of a table is the cheap way to do this, and
    # the right size for the job: twenty-nine million parameters against the
    # hundreds of millions of a document-layout pipeline, and it is asked only
    # where the table is, never what it says -- the words keep coming from the
    # text layer. Layout analysis is the fallback for an installation that does
    # not have it.
    if _tatr.available():
        return _read_shapes(path, pages, meta, pending)
    if not docling_available():
        return "\n\n".join(pages), meta, None

    read_by_model = 0
    with tempfile.TemporaryDirectory(prefix="mdcx-pages-") as carpeta:
        gathered = Path(carpeta) / "pending.pdf"
        try:
            # One document rather than one per page. Layout analysis charges a
            # fixed cost per document and batches the pages inside it, so four
            # pages sent separately cost more than the whole extract sent at
            # once -- measured, 21.7 seconds against 13.3.
            _pdf.extract_page_list(path, pending, gathered)
            md, _, _ = docling_convert(gathered, ocr=False)
        except Exception:  # noqa: BLE001
            md = ""             # the natively extracted pages stay as they are
        if md and md.strip():
            # The model reads the gathered pages as one document and does not
            # say where each began, so its reading is placed where the first of
            # them was and the others give up their native text to it. They are
            # pages whose table the cheap path could not produce, and their
            # prose is what the model returns.
            first = pending[0]
            pages[first] = f"\n<!-- pages {', '.join(str(i + 1) for i in pending)} -->\n\n{md.strip()}"
            for index in pending[1:]:
                pages[index] = ""
            read_by_model = len(pending)

    meta = dict(meta)
    meta["engine"] = "pypdfium2+docling" if read_by_model else "pypdfium2"
    meta["pages_read_by_model"] = read_by_model
    return "\n\n".join(p for p in pages if p), meta, None

def _md_escape_cell(value) -> str:
    if value is None:
        return ""
    text = str(value).replace("|", "\\|")
    return " ".join(text.split())

def native_xlsx(path: Path) -> tuple[str, dict, None]:
    """XLSX to one Markdown table per sheet."""
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out: list[str] = []
    sheets = []
    try:
        for ws in wb.worksheets:
            sheets.append(ws.title)
            out.append(f"\n## Hoja: {ws.title}\n")
            rows = [list(r) for r in ws.iter_rows(values_only=True)]
            while rows and all(c is None for c in rows[-1]):
                rows.pop()
            if not rows:
                out.append("_(hoja vacia)_")
                continue
            width = max(len(r) for r in rows)
            header = [_md_escape_cell(c) for c in rows[0]] + [""] * (width - len(rows[0]))
            out.append("| " + " | ".join(header) + " |")
            out.append("|" + "---|" * width)
            for row in rows[1:]:
                cells = [_md_escape_cell(c) for c in row] + [""] * (width - len(row))
                out.append("| " + " | ".join(cells) + " |")
    finally:
        wb.close()
    return "\n".join(out), {"engine": "openpyxl", "sheets": sheets}, None

def native_docx(path: Path) -> tuple[str, dict, None]:
    """DOCX -> Markdown respetando niveles de title, listas y tables."""
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    d = docx.Document(str(path))
    out: list[str] = []

    def render_paragraph(p) -> str:
        text = p.text.strip()
        if not text:
            return ""
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            m = re.search(r"(\d+)", style)
            level = min(int(m.group(1)), 6) if m else 1
            return ("#" * level) + " " + text
        if "list" in style:
            return "- " + text
        return text

    def render_table(t) -> list[str]:
        rows = [
            [" ".join(c.text.split()).replace("|", "\\|") for c in r.cells] for r in t.rows
        ]
        if not rows:
            return []
        width = max(len(r) for r in rows)
        lines = [
            "",
            "| " + " | ".join(rows[0] + [""] * (width - len(rows[0]))) + " |",
            "|" + "---|" * width,
        ]
        for r in rows[1:]:
            lines.append("| " + " | ".join(r + [""] * (width - len(r))) + " |")
        lines.append("")
        return lines

    body = d.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            line = render_paragraph(Paragraph(child, d))
            if line:
                out.append(line)
        elif tag == "tbl":
            out.extend(render_table(Table(child, d)))

    for section in d.sections:
        for name, container in (("Encabezado", section.header), ("Pie", section.footer)):
            texts = [p.text.strip() for p in container.paragraphs if p.text.strip()]
            if texts:
                out.append("\n<!-- " + name + ": " + " / ".join(texts) + " -->")

    return "\n\n".join(out), {"engine": "python-docx", "tables": len(d.tables)}, None

def native_text(path: Path) -> tuple[str, dict, None]:
    from .extract import _plain_text

    text, meta = _plain_text(path)
    meta["engine"] = "passthrough"
    return text, meta, None

def native_html(path: Path) -> tuple[str, dict, None]:
    raw = native_text(path)[0]
    raw = re.sub(r"(?is)<(script|style).*?</\1>", " ", raw)
    raw = re.sub(r"(?i)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</(p|div|tr|h[1-6]|li)>", "\n\n", raw)
    raw = re.sub(r"<[^>]+>", " ", raw)
    raw = html.unescape(re.sub(r"[ \t]{2,}", " ", raw))
    return raw, {"engine": "html-strip"}, None

NATIVE = {
    "pdf": native_pdf,
    "xlsx": native_xlsx,
    "docx": native_docx,
    "pptx": None,
    "text": native_text,
    "html": native_html,
}
