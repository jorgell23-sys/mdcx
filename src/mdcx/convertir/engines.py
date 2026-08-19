# Copyright 2026 Jorge Ellena G.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Motores de conversion a Markdown, en cascada de mayor a menor estructura.

Orden de preferencia:
  1. Docling (IBM, MIT): entiende layout, jerarquia de titulos y tablas (TableFormer).
     Da el Markdown mas legible y ademas un JSON sin perdida como respaldo.
  2. Extractores nativos deterministas (PyMuPDF / python-docx / openpyxl): no entienden
     estructura, pero no inventan nada y no dependen de modelos. Son la red de seguridad.

Ningun motor usa un LLM para transcribir: un modelo generativo puede parafrasear o
alucinar cifras, que es exactamente la perdida de informacion que se busca evitar.
El LLM local queda reservado para tareas de descripcion, no de transcripcion.
"""
from __future__ import annotations

import os as _os
import html
import re
from pathlib import Path

_CONVERTERS: dict = {}

# Aceleracion: se decide una vez por proceso. Sobre GPU los modelos de layout, TableFormer
# y OCR corren varias veces mas rapido; el resultado es el mismo (mismos pesos), pero al
# abaratar el costo permite dejar activadas en todos los documentos las opciones de mayor
# calidad (TableFormer ACCURATE y OCR de pagina completa) sin que el lote se vuelva inviable.
_DEVICE_CACHE: dict = {}


def gpu_available() -> bool:
    if "gpu" not in _DEVICE_CACHE:
        try:
            import torch

            _DEVICE_CACHE["gpu"] = bool(torch.cuda.is_available())
            _DEVICE_CACHE["name"] = (
                torch.cuda.get_device_name(0) if _DEVICE_CACHE["gpu"] else "CPU"
            )
        except Exception:
            _DEVICE_CACHE["gpu"] = False
            _DEVICE_CACHE["name"] = "CPU"
    return _DEVICE_CACHE["gpu"]


def device_name() -> str:
    gpu_available()
    return _DEVICE_CACHE.get("name", "CPU")


def _accelerator_options(threads: int = 4):
    """Configura Docling para usar CUDA cuando esta disponible, con caida limpia a CPU."""
    try:
        from docling.datamodel.pipeline_options import AcceleratorDevice, AcceleratorOptions

        device = AcceleratorDevice.CUDA if gpu_available() else AcceleratorDevice.AUTO
        return AcceleratorOptions(num_threads=threads, device=device)
    except Exception:
        return None


def local_artifacts_path():
    """Carpeta con los modelos ya descargados, si existe en este equipo.

    Se busca en la ubicacion estandar de la libreria y en una copia dentro del proyecto,
    para que la herramienta pueda distribuirse junto con sus pesos y funcionar en un
    equipo sin acceso a internet.
    """
    from pathlib import Path as _P
    import os as _os

    candidatas = []
    env = _os.environ.get("DOCLING_ARTIFACTS_PATH")
    if env:
        candidatas.append(_P(env))
    candidatas.append(_P(__file__).resolve().parent.parent / "models")
    candidatas.append(_P.home() / ".cache" / "docling" / "models")
    for c in candidatas:
        if c.is_dir() and any(c.iterdir()):
            return c
    return None


def docling_available() -> bool:
    try:
        import docling  # noqa: F401

        return True
    except Exception:
        return False


def _get_docling_converter(ocr: bool):
    """Crea (y cachea) un DocumentConverter. Cachear importa: construirlo carga modelos."""
    key = f"ocr={ocr}"
    if key in _CONVERTERS:
        return _CONVERTERS[key]

    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions

    popts = PdfPipelineOptions()
    popts.do_ocr = ocr
    popts.do_table_structure = True
    accel = _accelerator_options()
    if accel is not None:
        popts.accelerator_options = accel

    # Modelos desde el disco local. Sin esto, la libreria consulta el repositorio remoto
    # en cada arranque para comprobar versiones: la conversion pasaria a depender de la
    # red y de un servicio externo, cuando los pesos ya estan descargados. Fijar la ruta
    # deja el proceso completamente autonomo y reproducible.
    artefactos = local_artifacts_path()
    if artefactos is not None:
        popts.artifacts_path = str(artefactos)
    try:
        from docling.datamodel.pipeline_options import TableFormerMode

        # El modo rapido, y no el llamado exacto, por una razon medida y no por prisa.
        #
        # El proyecto uso ACCURATE desde el principio, dando por supuesto que su nombre
        # describia el resultado. Comparados los dos sobre los documentos con mas tablas del
        # corpus, resulto lo contrario de lo esperado:
        #
        #   Tiempo. FAST tarda un 40 % menos. En la lista de equipos mecanicos, 250 segundos
        #   frente a 421.
        #
        #   Contenido. Identico. Los dos modos terminan en el 100 % de cobertura, porque lo
        #   que el motor no estructura lo recoge el anexo de recuperacion. El modo de tabla
        #   no decide que texto se conserva; solo cuanto se tarda en llegar.
        #
        #   Estructura. FAST reconstruye igual o mejor: mismas 36 tablas, 285 filas frente a
        #   278 y 1.299 celdas con dato frente a 1.270. Y en el documento mas pesado captura
        #   17 puntos mas de texto que ACCURATE, que se degrada cuando la tabla es enorme.
        #
        # Se deja el modo exacto disponible en la variable de entorno PDFTOMD_TABLAS para
        # quien quiera contrastarlo sobre otro corpus: la medicion vale para estos documentos
        # y no es una ley general.
        modo = _os.environ.get("PDFTOMD_TABLAS", "rapido").strip().lower()
        popts.table_structure_options.mode = (
            TableFormerMode.ACCURATE if modo in ("exacto", "accurate")
            else TableFormerMode.FAST)
        # Alinea las celdas reconocidas con el texto real del PDF en vez de re-OCRizarlas.
        popts.table_structure_options.do_cell_matching = True
    except Exception:
        pass
    if ocr:
        # El reconocedor viene configurado para chino de fabrica. Sobre un plano en ingles
        # eso degrada visiblemente la lectura ("CONDENSATE" sale como "CONCEVSATE"), porque
        # el modelo de reconocimiento no comparte diccionario ni alfabeto. Se pide el
        # modelo latino, que es el que corresponde a estos documentos.
        try:
            from docling.datamodel.pipeline_options import RapidOcrOptions

            popts.ocr_options = RapidOcrOptions(
                lang=["english"],
                force_full_page_ocr=True,
                # Mas resolucion de la que trae por defecto: en un plano el texto es
                # pequeño frente al tamaño de la lamina y a escala baja se pierde.
                scale=4.0,
            )
        except Exception:
            try:
                popts.ocr_options.force_full_page_ocr = True
            except Exception:
                pass

    conv = DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=popts)}
    )
    _CONVERTERS[key] = conv
    return conv


def _full_export_kwargs() -> dict:
    """Fuerza a Docling a exportar TODO el documento, no solo el cuerpo.

    Por defecto el Markdown omite la capa `furniture` (encabezados, pies y numeros de
    pagina) y algunos tipos de elemento. En documentos de ingenieria esa capa contiene
    justamente el codigo de documento, la revision y el cliente, asi que excluirla es
    perdida de informacion. Se piden explicitamente todas las capas y todas las etiquetas.
    """
    kwargs: dict = {}
    try:
        from docling_core.types.doc.document import ContentLayer

        kwargs["included_content_layers"] = set(ContentLayer)
    except Exception:
        pass
    try:
        from docling_core.types.doc.labels import DocItemLabel

        kwargs["labels"] = set(DocItemLabel)
    except Exception:
        pass
    return kwargs


def docling_convert(path: Path, ocr: bool = False) -> tuple[str, dict, dict | None]:
    """Devuelve (markdown, metadatos, documento_json_sin_perdida)."""
    conv = _get_docling_converter(ocr)
    result = conv.convert(str(path))
    doc = result.document
    try:
        md = doc.export_to_markdown(**_full_export_kwargs())
    except Exception:
        md = doc.export_to_markdown()
    try:
        lossless = doc.export_to_dict()
    except Exception:
        lossless = None
    meta = {"engine": "docling", "ocr": ocr}
    try:
        meta["pages"] = len(doc.pages)
    except Exception:
        pass
    return md, meta, lossless


# --------------------------------------------------------------------------------------
# Motores nativos: sin modelos, sin dependencias pesadas, salida predecible.
# --------------------------------------------------------------------------------------


def native_pdf(path: Path) -> tuple[str, dict, None]:
    """PDF -> Markdown por bloques de layout. Preserva orden de lectura y separa paginas."""
    from . import pdf as _pdf

    doc = _pdf.leer_documento(path)
    out: list[str] = []
    try:
        for i, page in enumerate(doc, start=1):
            out.append(f"\n<!-- pagina {i} -->\n")
            # Se usa la via rapida: el motor nativo es la red de seguridad y su fidelidad
            # se mide contra el mismo texto plano que lee. Reconstruir columnas rectangulo a
            # rectangulo costaba un segundo por pagina y no cambiaba lo que se conserva.
            for text in _pdf.parrafos_rapidos(page):
                if text:
                    out.append(text)
        meta = {"engine": "pypdfium2", "pages": len(doc)}
    finally:
        doc.close()
    return "\n\n".join(out), meta, None


def _md_escape_cell(value) -> str:
    if value is None:
        return ""
    text = str(value).replace("|", "\\|")
    return " ".join(text.split())


def native_xlsx(path: Path) -> tuple[str, dict, None]:
    """XLSX -> una tabla Markdown por hoja, conservando celdas vacias y el nombre de hoja."""
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out: list[str] = []
    sheets = []
    try:
        for ws in wb.worksheets:
            sheets.append(ws.title)
            out.append(f"\n## Hoja: {ws.title}\n")
            rows = [list(r) for r in ws.iter_rows(values_only=True)]
            # Recortar filas totalmente vacias al final, pero no en el medio:
            # una fila vacia intermedia puede ser un separador con significado.
            while rows and all(c is None for c in rows[-1]):
                rows.pop()
            if not rows:
                out.append("_(hoja vacia)_")
                continue
            width = max(len(r) for r in rows)
            header = [_md_escape_cell(c) for c in rows[0]] + [""] * (width - len(rows[0]))
            out.append("| " + " | ".join(header) + " |")
            out.append("|" + "---|" * width)
            for row in rows[1:]:
                cells = [_md_escape_cell(c) for c in row] + [""] * (width - len(row))
                out.append("| " + " | ".join(cells) + " |")
    finally:
        wb.close()
    return "\n".join(out), {"engine": "openpyxl", "sheets": sheets}, None


def native_docx(path: Path) -> tuple[str, dict, None]:
    """DOCX -> Markdown respetando niveles de titulo, listas y tablas."""
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    d = docx.Document(str(path))
    out: list[str] = []

    def render_paragraph(p) -> str:
        text = p.text.strip()
        if not text:
            return ""
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            m = re.search(r"(\d+)", style)
            level = min(int(m.group(1)), 6) if m else 1
            return ("#" * level) + " " + text
        if "list" in style:
            return "- " + text
        return text

    def render_table(t) -> list[str]:
        rows = [
            [" ".join(c.text.split()).replace("|", "\\|") for c in r.cells] for r in t.rows
        ]
        if not rows:
            return []
        width = max(len(r) for r in rows)
        lines = [
            "",
            "| " + " | ".join(rows[0] + [""] * (width - len(rows[0]))) + " |",
            "|" + "---|" * width,
        ]
        for r in rows[1:]:
            lines.append("| " + " | ".join(r + [""] * (width - len(r))) + " |")
        lines.append("")
        return lines

    # Recorrer el cuerpo en el orden real del XML mantiene tablas y parrafos intercalados.
    body = d.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            line = render_paragraph(Paragraph(child, d))
            if line:
                out.append(line)
        elif tag == "tbl":
            out.extend(render_table(Table(child, d)))

    for section in d.sections:
        for name, container in (("Encabezado", section.header), ("Pie", section.footer)):
            texts = [p.text.strip() for p in container.paragraphs if p.text.strip()]
            if texts:
                out.append("\n<!-- " + name + ": " + " / ".join(texts) + " -->")

    return "\n\n".join(out), {"engine": "python-docx", "tables": len(d.tables)}, None


def native_text(path: Path) -> tuple[str, dict, None]:
    from .extract import _plain_text

    text, meta = _plain_text(path)
    meta["engine"] = "passthrough"
    return text, meta, None


def native_html(path: Path) -> tuple[str, dict, None]:
    raw = native_text(path)[0]
    raw = re.sub(r"(?is)<(script|style).*?</\1>", " ", raw)
    raw = re.sub(r"(?i)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</(p|div|tr|h[1-6]|li)>", "\n\n", raw)
    raw = re.sub(r"<[^>]+>", " ", raw)
    raw = html.unescape(re.sub(r"[ \t]{2,}", " ", raw))
    return raw, {"engine": "html-strip"}, None


NATIVE = {
    "pdf": native_pdf,
    "xlsx": native_xlsx,
    "docx": native_docx,
    "pptx": None,
    "text": native_text,
    "html": native_html,
}
